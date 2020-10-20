import json
from uuid import uuid1

from docx import Document
from docx.shared import Pt
from static.docx import *
import os

from utils.ClientModels import RunResult, Procedure, UsecaseGroup, Usecase


# result_dir = "D:\dcsdb\static"


def export_procedure_result(uid, run_time, result_dir, tmp_file=False):
    document = Document(docx = "./static/test_result_template.docx")
    run_result = RunResult.select().where(RunResult.run_uuid == uid)
    global result_uid
    global procedure
    result_uid = uid
    try:
        procedure = Procedure.get(Procedure.number == run_result.first().procedure_number)
    except:
        try:
            procedure = UsecaseGroup.get(UsecaseGroup.name == run_result.first().usecase_group_name)
        except:
            procedure = Usecase.get(Usecase.number == run_result.first().usecase_number)
    if type(procedure) == Usecase:
        allUseCase = [Usecase.number]
    else:
        allUseCase = json.loads(procedure.usecase)
    document.paragraphs[2].add_run("文件标题：" + procedure.name + "测试报告")
    document.paragraphs[2].runs[0].font.size = Pt(18)
    document.paragraphs[2].runs[0].bold = True
    first_table = document.tables[0]
    if type(procedure) == UsecaseGroup:
        first_table.cell(1, 1).text = " " * 17 + procedure.usecase_group_number
    else:
        first_table.cell(1, 1).text = " " * 17 + procedure.number
    first_table.cell(1, 1).paragraphs[0].runs[0].font.bold = True
    title_number = ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7", "5.8", "5.9", "5.10",
                    "5.11", "5.12", "5.13", "5.14", "5.15", "5.16", "5.17", "5.18", "5.19", "5.20", ]
    for index ,usecase_number in enumerate(allUseCase):
        title_name = title_number[index]
        export_usecase_result(document, title_name, usecase_number, run_time, result_dir)
    newname = procedure.name + str(run_time).replace(":", "") + "测试报告.doc"
    if tmp_file:
        newname = str(uuid1()) + ".doc"
    import os
    file_save_path = os.path.join(result_dir, newname)

    document.save(file_save_path)
    # if tmp_file:
    #     return word2pdf(file_save_path)
    return file_save_path


def export_usecase_result(document, title_number, usecase_number, run_time, result_dir):
    usecase = Usecase.get(Usecase.number == usecase_number)
    document.add_heading(text=title_number + usecase.name, level=2).runs[0].font.size = Pt(12)
    table = document.add_table(rows=3, cols=9, style='Table Grid')
    table.cell(0, 0).text = "测试用例"
    # table.cell(0, 0)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 1).text = usecase.name
    table.cell(0, 2).text = "用例编号"
    # table.cell(0, 2)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 3).text = usecase_number
    table.cell(0, 4).text = "工况描述"
    # table.cell(0, 4)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
    table.cell(0, 5).merge(table.cell(0, 6)).text = usecase.description
    # table.cell(0, 5)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
    table.cell(0, 7).text = "IC"
    # table.cell(0, 7)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 8).text = usecase.IC

    table.cell(1, 0).text = "序号"
    table.cell(1, 1).merge(table.cell(1, 2)).text = "实验步骤"
    table.cell(1, 3).text = "预期结果"
    table.cell(1, 4).text = "实际结果"
    table.cell(1, 5).merge(table.cell(1, 6)).text = "实际与预期一致"
    table.cell(1, 7).text = "测试时间"
    table.cell(1, 8).text = "备注"

    # for sub_cell in table.row_cells(1):
    #     sub_cell._tc.get_or_add_tcPr().append(
    #         parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(2, 1).text = "操作"
    table.cell(2, 2).text = "位置"
    table.cell(2, 5).text = "是"
    table.cell(2, 6).text = "否"
    table.autofit = False
    export_usecase_operation(table, usecase.number, json.loads(usecase.operation))


def export_usecase_operation(table, usecase_number, operation):
    for rsection in operation:
        export_usecase_section(table, usecase_number, rsection)


def export_usecase_section(table, usecase_number, rsection):
    addrow = table.add_row().cells
    for sub_cell in addrow[1:len(addrow)]:
        addrow[0].merge(sub_cell)
    addrow[0].text = rsection[0]
    runresult = RunResult.select().where(
                                         RunResult.usecase_number == usecase_number).first()
    run_text = json.loads(runresult.run_text)

    for i in range(1, len(rsection)):
        for j in range(1, len(rsection[i])):
            sub_add_row = table.add_row().cells
            sub_add_row[0].text = rsection[i][j]["sort"]
            sub_add_row[1].text = rsection[i][j]["name"]
            sub_add_row[2].text = rsection[i][j]["location"]
            # try:
            if rsection[i][0] == "WRITE":
                # sub_add_row[3].text = rsection[i][j]["except"]
                sub_add_row[4].text = run_text[rsection[i][j]["name"]][-1]
                if run_text[rsection[i][j]["name"]][0] == '是':
                    sub_add_row[5].text = '是'
                else:
                    sub_add_row[6].text = '否'
                sub_add_row[7].text = run_text[rsection[i][j]["name"]][1]

            if rsection[i][0] == "READ":
                # print(rsection[i][j])
                sub_add_row[3].text = rsection[i][j]["except"]
                sub_add_row[4].text = run_text[rsection[i][j]["name"]][-1]
                if run_text[rsection[i][j]["name"]][0] == '是':
                    sub_add_row[5].text = '是'
                else:
                    sub_add_row[6].text = '否'
                sub_add_row[7].text = run_text[rsection[i][j]["name"]][1]
            # if rsection[i][0] != "WRITE":
            #     sub_add_row[3].text = rsection[i][j]["except"]
            # sub_add_row[2]._tc.get_or_add_tcPr().append(
            #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
            # if rsection[i][0] != "WRITE":
            #     sub_add_row[3].text = rsection[i][j]["except"]
            # sub_add_row[3]._tc.get_or_add_tcPr().append(
            #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
                                                 # RunResult.section_sort == int(rsection[i][j]["sort"])).first()
            # if runresult:
            #     # sub_add_row[4].text = run_text  # 实际结果
            #     if runresult.run_result:
            #         sub_add_row[5].text = "是"  # 是
            #         # sub_add_row[5]._tc.get_or_add_tcPr().append(
            #         #     parse_xml(r'<w:shd {} w:fill="33CC52"/>'.format(nsdecls('w'))))
            #     else:
            #         sub_add_row[6].text = "否"  # 否
                    # sub_add_row[6]._tc.get_or_add_tcPr().append(
                    #     parse_xml(r'<w:shd {} w:fill="EE3D11"/>'.format(nsdecls('w'))))
                # sub_add_row[7].text = str(runresult.run_time)  # 测试时间
            sub_add_row[8].text = str(rsection[i][j]["remark"])  # 备注
            # except:
            #     print('error')
            #     continue


def exportCertification(uid, run_time, result_dir):
    print(result_dir)
    run_result = RunResult.select().where(RunResult.run_uuid == uid).first()
    if run_result.run_type == 1:
        name = run_result.procedure_name
    if run_result.run_type == 2:
        name = run_result.usecase_group_name
    if run_result.run_type == 3:
        name = run_result.usecase_number

    certification = json.loads(run_result.certification)

    document = Document()
    document.add_heading(name + '自证报告', 0)  #插入标题

    for k, v in certification.items():
        action = k.replace('\n', '') + '\n'
        p = document.add_paragraph(action)   #插入段落
        if v[0] == 'SET':
            if len(v) == 3:
                p.add_run(f'解析{action},通过短语库字段检索,判定该步骤为{v[0]},解析到需要操作的变量为{v[1]},设置失败,错误原因{v[2]}')
            elif len(v) == 4:
                p.add_run(f'解析{action},通过短语库字段检索,判定该步骤为{v[0]},解析到需要操作的变量为{v[1]},通过{v[2][0]}通道,使用{v[2][1]}协议,将变量{v[1]}的值定义为{v[3]}')
        elif v[0] == 'CHECK':
            if len(v) == 3:
                p.add_run(f'解析{action},通过短语库字段检索,判定该步骤为{v[0]},解析到需要操作的变量为{v[1]},读取失败,错误原因为{v[2]}')
            elif len(v) == 5:
                if v[-1]:
                    p.add_run(f'解析{action},通过短语库字段检索,判定该步骤为{v[0]},解析到需要操作的变量为{v[1]},通过{v[2][0]}通道,使用{v[2][1]}协议,检测到变量{v[1]}的值为{v[-2]},与预期结果{action}相同')
                else:
                    p.add_run(f'解析{action},通过短语库字段检索,判定该步骤为{v[0]},解析到需要操作的变量为{v[1]},通过{v[2][0]}通道,使用{v[2][1]}协议,检测到变量{v[1]}的值为{v[-2]},与预期结果{action}不符')
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    document.save(os.path.join(result_dir, name + '自证.docx'))



# if phrase == 'SET':
#     return f'通过{channel}通道,使用{protocol}协议,将变量{name}的值定义为{value}'
# if phrase == 'CHECK':
#     return f'通过{channel}通道,使用{protocol}协议,检查变量{name}的值为{value}'