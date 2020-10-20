# -*- coding: UTF-8 -*-
"""
    导出规程的测试结果，根据运行的uid直接获取本次的测试结果

"""
from procedure.manage_procedure.models import RunResult, Procedure, Usecase, UsecaseGroup
from ui_dcs.ui_utils import GetStaticPath
from docx import Document
from docx.shared import Pt
from uuid import uuid1
import json


def export_procedure_result(uid, run_time, result_dir, tmp_file=False):
    # 导出规程报告
    document = Document(docx=GetStaticPath("test_result_template.docx"))
    global result_uid
    global procedure
    result_uid = uid
    run_result = RunResult.select().where(RunResult.run_uuid == uid)
    try:
        procedure = Procedure.get(Procedure.number == run_result.first().procedure_number)
    except:
        procedure = UsecaseGroup.get(UsecaseGroup.name == run_result.first().usecase_group_name)
    document.paragraphs[2].add_run(u"文件标题：" + procedure.name + u"测试报告")
    document.paragraphs[2].runs[0].font.size = Pt(18)
    document.paragraphs[2].runs[0].bold = True
    first_table =  document.tables[0]
    first_table.cell(1,1).text = " "*17+procedure.number
    first_table.cell(1, 1).paragraphs[0].runs[0].font.bold = True

    # for section in document.sections:
    #     print ("+++++++",section.header)
    title_number = ["5.1", "5.2", "5.3", "5.4", "5.5", "5.6", "5.7", "5.8", "5.9", "5.10",
                    "5.11","5.12","5.13","5.14","5.15","5.16","5.17","5.18","5.19","5.20",]
    for usecase_number in json.loads(procedure.usecase):
        title_name = title_number[json.loads(procedure.usecase).index(usecase_number)]
        export_usecase_result(document, title_name, usecase_number, run_time, result_dir)
    newname = procedure.name + str(run_time).replace(":", "") + u"测试报告.doc"
    if tmp_file:
        newname = str(uuid1()) + ".doc"
    import os
    file_save_path = os.path.join(result_dir, newname)

    document.save(file_save_path)
    # if tmp_file:
    #     return word2pdf(file_save_path)
    return file_save_path


def export_usecase_result(document, title_number, usecase_number, run_time, result_dir):
    # 导出用例报告
    usecase = Usecase.get(Usecase.number == usecase_number)
    document.add_heading(text=title_number + usecase.name, level=2).runs[0].font.size = Pt(12)
    table = document.add_table(rows=3, cols=9, style='Table Grid')
    table.cell(0, 0).text = u"测试用例"
    # table.cell(0, 0)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 1).text = usecase.name
    table.cell(0, 2).text = u"用例编号"
    # table.cell(0, 2)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 3).text = usecase_number
    table.cell(0, 4).text = u"工况描述"
    # table.cell(0, 4)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
    table.cell(0, 5).merge(table.cell(0, 6)).text = usecase.description
    # table.cell(0, 5)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
    table.cell(0, 7).text = "IC"
    # table.cell(0, 7)._tc.get_or_add_tcPr().append(
    #     parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(0, 8).text = usecase.IC

    table.cell(1, 0).text = u"序号"
    table.cell(1, 1).merge(table.cell(1, 2)).text = u"实验步骤"
    table.cell(1, 3).text = u"预期结果"
    table.cell(1, 4).text = u"实际结果"
    table.cell(1, 5).merge(table.cell(1, 6)).text = u"实际与预期一致"
    table.cell(1, 7).text = u"测试时间"
    table.cell(1, 8).text = u"备注"

    # for sub_cell in table.row_cells(1):
    #     sub_cell._tc.get_or_add_tcPr().append(
    #         parse_xml(r'<w:shd {} w:fill="CC9C56"/>'.format(nsdecls('w'))))
    table.cell(2, 1).text = u"操作"
    table.cell(2, 2).text = u"位置"
    table.cell(2, 5).text = u"是"
    table.cell(2, 6).text = u"否"
    table.autofit = False
    export_usecase_operation(table, usecase.number, json.loads(usecase.operation))
    # TODO 设置页眉和页脚


def export_usecase_operation(table, usecase_number, operation):
    # 导出用例步骤
    for rsection in operation:
        export_usecase_section(table, usecase_number, rsection)


def export_usecase_section(table, usecase_number, rsection):
    # 导出用例
    addrow = table.add_row().cells
    for sub_cell in addrow[1:len(addrow)]:
        addrow[0].merge(sub_cell)
    addrow[0].text = rsection[0]

    for i in range(1, len(rsection)):
        for j in range(1, len(rsection[i])):
            sub_add_row = table.add_row().cells
            sub_add_row[0].text = rsection[i][j]["sort"]
            sub_add_row[1].text = rsection[i][j]["name"]
            sub_add_row[2].text = rsection[i][j]["location"]
            if rsection[i][0] != "WRITE":
                sub_add_row[3].text = rsection[i][j]["except"]
            if rsection[i][0] != "WRITE":
                sub_add_row[3].text = rsection[i][j]["except"]
            # sub_add_row[2]._tc.get_or_add_tcPr().append(
            #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
            if rsection[i][0] != "WRITE":
                sub_add_row[3].text = rsection[i][j]["except"]
            # sub_add_row[3]._tc.get_or_add_tcPr().append(
            #     parse_xml(r'<w:shd {} w:fill="EEEE11"/>'.format(nsdecls('w'))))
            runresult = RunResult.select().where(RunResult.run_uuid == result_uid,
                                                 RunResult.usecase_number == usecase_number,
                                                 RunResult.section_sort == int(rsection[i][j]["sort"])).first()
            if runresult:
                sub_add_row[4].text = runresult.run_text  # 实际结果
                if runresult.run_result:
                    sub_add_row[5].text = u"是"  # 是
                    # sub_add_row[5]._tc.get_or_add_tcPr().append(
                    #     parse_xml(r'<w:shd {} w:fill="33CC52"/>'.format(nsdecls('w'))))
                else:
                    sub_add_row[6].text = u"否"  # 否
                    # sub_add_row[6]._tc.get_or_add_tcPr().append(
                    #     parse_xml(r'<w:shd {} w:fill="EE3D11"/>'.format(nsdecls('w'))))
                sub_add_row[7].text = str(runresult.run_time)  # 测试时间
                sub_add_row[8].text = rsection[i][j]["remark"]  # 备注

import sys, os
from win32com.client import Dispatch, constants, gencache

def word2pdf(filename):

  input=filename
  output=input.split(".")[0]+".pdf"
  pdf_name=output

  #判断文件是否存在
  # os.chdir(input)
  if not os.path.isfile(input):
    print u'%s not exist'%input
    return False
  #文档路径需要为绝对路径，因为Word启动后当前路径不是调用脚本时的当前路径。
  if (not os.path.isabs(input)):#判断是否为绝对路径
    #os.chdir(REPORT_DOC_PATH)
    input = os.path.abspath(input)#返回绝对路径
  else:
    print u'%s not absolute path'%input
    return False
  if (not os.path.isabs(output)):
    os.chdir(output)
    output = os.path.abspath(output)
  else:
    print u'%s not absolute path'%output
    return False

  try:
    print input,output
    # enable python COM support for Word 2007
    # this is generated by: makepy.py -i "Microsoft Word 12.0 Object Library"
    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
    #开始转换
    w = Dispatch("Word.Application")
    try:
      doc = w.Documents.Open(input, ReadOnly = 1)
      doc.ExportAsFixedFormat(output, constants.wdExportFormatPDF,\
        Item = constants.wdExportDocumentWithMarkup, CreateBookmarks = constants.wdExportCreateHeadingBookmarks)
    except:
      print ' exception'
    finally:
      w.Quit(constants.wdDoNotSaveChanges)

    if os.path.isfile(pdf_name):
      print 'translate success'
      return pdf_name
    else:
      print 'translate fail'
      return False
  except:
    print ' exception'
    return -1

